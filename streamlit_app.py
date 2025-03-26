import pandas as pd
import json
import streamlit as st
from concurrent.futures import ThreadPoolExecutor
from openai import OpenAI
import docx
import openpyxl
import email
import os
import re
import tempfile
from io import BytesIO
import base64
import sys

# Import the PDF processing function from markdown.py
sys.path.append(os.path.dirname(os.path.abspath(__file__)))
from markdown import process_pdf

# Initialize OpenAI client
client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])

# Modified System prompt for extracting multiple transportation orders
SYSTEM_PROMPT = """
You are an expert data extraction system specialized in analyzing transportation and logistics documents. Your task is to carefully extract information from purchase orders, invoices, and transport documents.

IMPORTANT: A single document may contain MULTIPLE transportation orders or trips. You must identify and extract ALL separate transportation orders in the document and return them as an ARRAY of JSON objects.

For each transportation order, extract the following fields EXACTLY as they appear in the document or as close as possible:
- FROM_LOCATION: The origin location of transportation
- TO_LOCATION: The destination location
- MATERIAL_TYPE: Type of materials being transported (e.g., "Autoparts", "Steel", "Electronics")
- BODY_TYPE: Vehicle or container type (e.g., "20FT Container", "Truck", "Van")
- FREQUENCY: How frequent the shipment occurs (numeric value if possible)
- WEIGHT: Weight of goods with unit (e.g., "200 T", "500 kg")
- RATE_UOM: Rate unit of measure (e.g., "Per KG", "Per Ton", "Per Trip")
- TRIPS_IN_MONTH: Number of trips per month (numeric value)
- START_DATE: Contract start date in YYYY-MM-DD format
- END_DATE: Contract end date in YYYY-MM-DD format

For any field not found in the document, use null. Analyze the entire document carefully, looking for these fields in headers, tables, key-value pairs, and body text.

Look carefully for cases where:
1. The document contains multiple FROM_LOCATION and TO_LOCATION pairs
2. Same FROM_LOCATION and TO_LOCATION but with different MATERIAL_TYPE or BODY_TYPE
3. Tables with rows representing different transportation orders

Return ONLY a valid JSON array of objects with these fields and nothing else.
Each object in the array represents a separate transportation order.

Your output must strictly follow this JSON structure:
[
  {
    "FROM_LOCATION": "Location1",
    "TO_LOCATION": "Location2",
    "MATERIAL_TYPE": "MaterialType1",
    "BODY_TYPE": "BodyType1",
    "FREQUENCY": "Frequency1",
    "WEIGHT": "Weight1",
    "RATE_UOM": "RateUOM1",
    "TRIPS_IN_MONTH": "TripsPerMonth1",
    "START_DATE": "StartDate1",
    "END_DATE": "EndDate1"
  },
  {
    "FROM_LOCATION": "Location3",
    "TO_LOCATION": "Location4",
    "MATERIAL_TYPE": "MaterialType2",
    "BODY_TYPE": "BodyType2",
    "FREQUENCY": "Frequency2",
    "WEIGHT": "Weight2",
    "RATE_UOM": "RateUOM2",
    "TRIPS_IN_MONTH": "TripsPerMonth2",
    "START_DATE": "StartDate2",
    "END_DATE": "EndDate2"
  }
]

If you only find one transportation order, still return it as an array with one object.
"""

# Modified User prompt template
USER_PROMPT_TEMPLATE = """
Extract ALL transportation order details from the following document text. 
The document may contain multiple transportation orders/trips.
Return ONLY a JSON array following this exact structure:
[
  {
    "FROM_LOCATION": "Ambattur",
    "TO_LOCATION": "Hyderabad",
    "MATERIAL_TYPE": "Autoparts",
    "BODY_TYPE": "20FT Container",
    "FREQUENCY": "10",
    "WEIGHT": "200 T",
    "RATE_UOM": "Per KG",
    "TRIPS_IN_MONTH": "4",
    "START_DATE": "2019-11-06",
    "END_DATE": "2022-11-06"
  },
  {
    "FROM_LOCATION": "Location2",
    "TO_LOCATION": "Location2",
    "MATERIAL_TYPE": "Material2",
    "BODY_TYPE": "Container2",
    "FREQUENCY": "Frequency2",
    "WEIGHT": "Weight2",
    "RATE_UOM": "RateUOM2",
    "TRIPS_IN_MONTH": "Trips2",
    "START_DATE": "StartDate2",
    "END_DATE": "EndDate2"
  }
]

Here's the document content:
"""

# Function to extract text from DOCX files
def process_docx(file_content):
    markdown_text = "## Document data\n\n"
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as temp_file:
        temp_file.write(file_content)
        temp_path = temp_file.name
    
    try:
        doc = docx.Document(temp_path)
        
        # Extract paragraphs
        markdown_text += "\n### Text Content\n"
        for para in doc.paragraphs:
            if para.text.strip():
                markdown_text += para.text + "\n\n"
        
        # Extract tables
        if doc.tables:
            markdown_text += "\n### Tables\n"
            for i, table in enumerate(doc.tables):
                markdown_text += f"\n#### Table {i+1}\n"
                
                # Create the markdown table
                for row in table.rows:
                    cells = [cell.text.replace("\n", " ").strip() for cell in row.cells]
                    markdown_text += "| " + " | ".join(cells) + " |\n"
                    
                    # Add separator row after header (first row)
                    if row == table.rows[0]:
                        markdown_text += "| " + " | ".join(["---"] * len(row.cells)) + " |\n"
        
        return markdown_text
    finally:
        os.unlink(temp_path)

# Function to extract text from Excel files with improved debugging
def process_xlsx(file_content):
    markdown_text = "## Document data\n\n"
    
    with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as temp_file:
        temp_file.write(file_content)
        temp_path = temp_file.name
    
    try:
        # Load the workbook from file
        workbook = openpyxl.load_workbook(temp_path)
        
        # Process each worksheet
        for sheet_name in workbook.sheetnames:
            sheet = workbook[sheet_name]
            markdown_text += f"\n### Sheet: {sheet_name}\n\n"
            
            # Find the used range
            min_row, min_col = 1, 1
            max_row = max(1, sheet.max_row)
            max_col = max(1, sheet.max_column)
            
            # Debug info
            markdown_text += f"*Sheet dimensions: {min_row}:{max_row} rows, {min_col}:{max_col} columns*\n\n"
            
            # Skip empty sheets
            if max_row <= 1 and max_col <= 1 and not sheet.cell(1, 1).value:
                markdown_text += "*Empty sheet*\n\n"
                continue
                
            # Create header row for markdown table
            header_row = []
            for col in range(min_col, max_col + 1):
                col_letter = openpyxl.utils.get_column_letter(col)
                header_row.append(f"Col {col_letter}")
            
            markdown_text += "| " + " | ".join(header_row) + " |\n"
            markdown_text += "| " + " | ".join(["---"] * len(header_row)) + " |\n"
            
            # Add data rows
            for row in range(min_row, max_row + 1):
                row_data = []
                for col in range(min_col, max_col + 1):
                    cell_value = sheet.cell(row=row, column=col).value
                    row_data.append(str(cell_value) if cell_value is not None else "")
                
                markdown_text += "| " + " | ".join(row_data) + " |\n"
        
        return markdown_text
    except Exception as e:
        return f"## Error processing Excel file\n\nError: {str(e)}\n\nPlease ensure the file is a valid Excel file."
    finally:
        os.unlink(temp_path)

# Function to extract text from email (.eml) files
def process_eml(file_content):
    markdown_text = "## Email data\n\n"
    
    try:
        # Parse the email content
        msg = email.message_from_bytes(file_content)
        
        # Extract headers
        markdown_text += "### Email Headers\n"
        for header in ['From', 'To', 'Subject', 'Date']:
            value = msg.get(header, "")
            if value:
                markdown_text += f"- **{header}**: {value}\n"
        
        # Extract body
        markdown_text += "\n### Email Body\n"
        
        if msg.is_multipart():
            for part in msg.walk():
                content_type = part.get_content_type()
                content_disposition = str(part.get("Content-Disposition"))
                
                # Extract text parts
                if content_type == "text/plain" and "attachment" not in content_disposition:
                    try:
                        body = part.get_payload(decode=True).decode()
                        markdown_text += body + "\n\n"
                    except:
                        markdown_text += "(Unable to decode plain text body)\n\n"
                elif content_type == "text/html" and "attachment" not in content_disposition:
                    try:
                        html_body = part.get_payload(decode=True).decode()
                        # Very simple HTML to text conversion
                        text_body = re.sub('<[^<]+?>', ' ', html_body)
                        markdown_text += text_body + "\n\n"
                    except:
                        markdown_text += "(Unable to decode HTML body)\n\n"
        else:
            try:
                body = msg.get_payload(decode=True).decode()
                markdown_text += body + "\n\n"
            except:
                markdown_text += "(Unable to decode email body)\n\n"
        
        # List attachments
        markdown_text += "\n### Attachments\n"
        has_attachments = False
        
        for part in msg.walk():
            if part.get_content_maintype() == 'multipart':
                continue
            if part.get('Content-Disposition') is None:
                continue
            
            filename = part.get_filename()
            if filename:
                has_attachments = True
                markdown_text += f"- {filename}\n"
        
        if not has_attachments:
            markdown_text += "No attachments found.\n"
        
        return markdown_text
    except Exception as e:
        return f"## Error processing Email file\n\nError: {str(e)}"

# Function to process a file based on its type
def process_file(file_content, file_extension):
    # Create a temporary file
    with tempfile.NamedTemporaryFile(delete=False, suffix=file_extension) as temp_file:
        temp_file.write(file_content)
        temp_path = temp_file.name
    
    try:
        if file_extension.lower() == '.pdf':
            # Use the imported process_pdf function from markdown.py
            return process_pdf(temp_path)
        elif file_extension.lower() == '.docx':
            return process_docx(file_content)
        elif file_extension.lower() in ['.xlsx', '.xls']:
            return process_xlsx(file_content)
        elif file_extension.lower() == '.eml':
            return process_eml(file_content)
        else:
            return f"Unsupported file type: {file_extension}"
    finally:
        # Clean up the temporary file
        os.unlink(temp_path)

# Modified function to process a row from the dataframe
def process_row(row):
    file_name = row['file_name']
    file_content = row['file_content']
    file_extension = os.path.splitext(file_name)[1]
    
    # Ensure file_content is bytes
    if isinstance(file_content, str):
        try:
            # Try to decode if it's base64 encoded
            file_content = base64.b64decode(file_content)
        except:
            # Otherwise just encode the string
            file_content = file_content.encode('utf-8')
    
    # Extract text from the file based on its type
    extracted_text = process_file(file_content, file_extension)
    
    # Create the user prompt by combining the template and extracted text
    user_prompt = USER_PROMPT_TEMPLATE + "\n" + extracted_text
    
    try:
        # Call the LLM
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": SYSTEM_PROMPT},
                {"role": "user", "content": user_prompt}
            ],
            temperature=0.2,  # Even lower temperature for more consistent extraction
            max_tokens=2000,  # Further increased token limit to handle multiple orders
            top_p=1,
            frequency_penalty=0,
            presence_penalty=0,
            response_format={"type": "json_object"}  # Force JSON format response
        )
        
        # Extract the content from the response
        content = response.choices[0].message.content
        
        # Remove code block markers if present
        if content.startswith("```json"):
            content = content[7:]
        if content.endswith("```"):
            content = content[:-3]
        
        # Parse the cleaned content into JSON
        content = content.strip()
        
        try:
            # First try parsing the content directly
            response_json = json.loads(content)
        except json.JSONDecodeError:
            # If parsing fails, try extracting the JSON array from the text
            # This handles cases where the model might add additional text
            json_pattern = r'\[\s*{[\s\S]*}\s*\]'
            json_match = re.search(json_pattern, content)
            
            if json_match:
                json_str = json_match.group(0)
                response_json = json.loads(json_str)
            else:
                # If all else fails, create an empty list
                print(f"Failed to parse JSON from response: {content[:100]}...")
                response_json = []
        
        # Ensure response is a list/array
        if not isinstance(response_json, list):
            response_json = [response_json]  # Convert to list if single object
        
        # Required fields
        required_fields = [
            "FROM_LOCATION", "TO_LOCATION", "MATERIAL_TYPE", "BODY_TYPE", 
            "FREQUENCY", "WEIGHT", "RATE_UOM", "TRIPS_IN_MONTH", 
            "START_DATE", "END_DATE"
        ]
        
        # Create list to hold normalized data for each transportation order
        all_normalized = []
        
        # Process each transportation order
        for order_json in response_json:
            # Ensure all required fields are present
            for field in required_fields:
                if field not in order_json:
                    order_json[field] = None
            
            # Normalize the JSON into tabular format
            normalized = pd.json_normalize(order_json)
            
            # Add the file name for reference
            normalized["file_name"] = file_name
            
            all_normalized.append(normalized)
        
        # Combine all orders from this file
        if all_normalized:
            return pd.concat(all_normalized, ignore_index=True)
        else:
            return None
    
    except Exception as e:
        print(f"Error processing file {file_name}: {e}")
        # Create a default dataframe with null values
        default_data = {
            "FROM_LOCATION": None, "TO_LOCATION": None, "MATERIAL_TYPE": None,
            "BODY_TYPE": None, "FREQUENCY": None, "WEIGHT": None,
            "RATE_UOM": None, "TRIPS_IN_MONTH": None, "START_DATE": None,
            "END_DATE": None, "file_name": file_name
        }
        return pd.DataFrame([default_data])

# Main function to process all files
def process_files(input_file, output_file, max_workers=10):
    # Read the input dataframe
    df = pd.read_csv(input_file)
    
    # Ensure the required columns exist
    if not {'file_name', 'file_content'}.issubset(df.columns):
        raise ValueError("Input CSV must contain 'file_name' and 'file_content' columns.")
        
    # If the file_content is stored as a string in the CSV, we need to convert it back to bytes
    if isinstance(df['file_content'].iloc[0], str):
        try:
            # If it's stored as a base64 string
            df['file_content'] = df['file_content'].apply(lambda x: base64.b64decode(x) if isinstance(x, str) else x)
        except:
            # If it's just a regular string
            df['file_content'] = df['file_content'].apply(lambda x: x.encode('utf-8') if isinstance(x, str) else x)
    
    # Placeholder for normalized data
    normalized_data = []
    
    # Use ThreadPoolExecutor for parallel processing
    with ThreadPoolExecutor(max_workers=max_workers) as executor:
        futures = [executor.submit(process_row, row) for _, row in df.iterrows()]
        
        for future in futures:
            result = future.result()
            if result is not None:
                normalized_data.append(result)
    
    # Combine all normalized data into a single DataFrame
    if normalized_data:
        final_df = pd.concat(normalized_data, ignore_index=True)
        
        # Save the DataFrame to the output file
        final_df.to_csv(output_file, index=False)
        print(f"Normalized data saved to {output_file}")
        return final_df
    else:
        print("No data to save.")
        return None

# Streamlit App Example
def streamlit_app():
    st.title("Transportation Order Data Extractor")
    
    st.write("Upload transportation order files (.pdf, .docx, .xlsx, .eml) to extract structured data")
    
    uploaded_files = st.file_uploader("Choose files", accept_multiple_files=True, 
                                     type=["pdf", "docx", "xlsx", "xls", "eml"])
    
    if uploaded_files and st.button("Process Files"):
        # Create temporary dataframe to store files
        data = []
        
        for file in uploaded_files:
            # Read as bytes
            file_content = file.read()
            # Ensure we're dealing with bytes
            if not isinstance(file_content, bytes):
                file_content = file_content.encode('utf-8')
            data.append({"file_name": file.name, "file_content": file_content})
        
        temp_df = pd.DataFrame(data)
        
        # Save temporary CSV
        temp_input = "temp_input.csv"
        temp_output = "temp_output.csv"
        
        # Save file_content as base64 to ensure proper storage in CSV
        temp_df['file_content'] = temp_df['file_content'].apply(lambda x: base64.b64encode(x).decode('utf-8'))
        temp_df.to_csv(temp_input, index=False)
        
        with st.spinner("Processing files..."):
            results_df = process_files(temp_input, temp_output)
            
        if results_df is not None:
            st.success("Processing complete!")
            st.dataframe(results_df)
            
            # Allow user to download results
            csv = results_df.to_csv(index=False)
            st.download_button(
                label="Download CSV Results",
                data=csv,
                file_name="transportation_orders_data.csv",
                mime="text/csv"
            )
            
            # Also offer JSON format
            json_data = results_df.to_json(orient="records")
            st.download_button(
                label="Download JSON Results",
                data=json_data,
                file_name="transportation_orders_data.json",
                mime="application/json"
            )
        else:
            st.error("Error processing files")
            
# If running directly, start the Streamlit app
if __name__ == "__main__":
    streamlit_app()